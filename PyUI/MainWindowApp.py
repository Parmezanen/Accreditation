
"""!!!Файл взаимодействия главного окна с дочерними окнами!!!"""


import sys
import re
import os
import comtypes.client
import csv
import time
import docx
import pandas as pd
import numpy as np
from docx.shared import Pt
from docx.shared import Inches
#from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from PyQt5 import QtCore, QtGui, QtWidgets

"""Импорт файлов интерфейса"""
from UI.PPSEditor import Ui_PPSReference
from UI.MainWindow import Ui_MainWindow
from UI.UPDB_Editor import *
from UI.AudienceDB_Edit import *
from UI.KODB_Editor import *

from Sorting import SelSortAud
from Sorting import SelSortPPS

from SaveAndLoad import writeCSV
from SaveAndLoad import PPSreadCSV
from SaveAndLoad import AUDreadCSV
from SaveAndLoad import UPreadCSV
from SaveAndLoad import TeacherreadCSV
from SaveAndLoad import readMTODisc
from SaveAndLoad import readMTOAUD
from SaveAndLoad import matchingMTO
from SaveAndLoad import readKODisc
from SaveAndLoad import readKOTeacher
from SaveAndLoad import matchingKO


#Импорт функций генерации документов
import DocxGeneratingDef
#Импорт функции для конвертации файла
import convertDocxToPDF
#Импорт валидатора
import Validator

"""Инициализация классов интерфейса для их вызова в приложении"""


#Класс редактора аудиторий
class AudienceEditorWindow(QtWidgets.QMainWindow):
    def __init__(self):
        super(AudienceEditorWindow, self).__init__()
        self.records = []
        self.record={}
        self.index = 0
        self.Aaud = ""

        self.counter=0

        self.ui=Ui_Audience_Editor()


        self.ANDialogUi=AudNameDialog()
        self.ANDialogUi.setupUi(self)

        self.ANMDialogUi=AudNaimDialog()
        self.ANMDialogUi.setupUi(self)

        self.ATDialogUi=AudTODialog()
        self.ATDialogUi.setupUi(self)

        self.APDialogUi=AudPODialog()
        self.APDialogUi.setupUi(self)


        self.ui.setupUi(self)

        self.tableRecords()
        self.ui.pb_Save.clicked.connect(self.saveRecord)
        self.ui.pb_Add.clicked.connect(self.addRecord)
       # self.ui.pb_Add.clicked.connect(self.validation)
        self.ui.pb_Delete.clicked.connect(self.delRecord)
        self.ui.pb_Edit.clicked.connect(self.editRecord)
        self.ui.pb_Save.setEnabled(False)

        self.ui.tb_Audience.currentCellChanged.connect(self.ShowRecord)
        #if self.ui.tb_Audience.cellClicked(self.ui.tb_Audience.currentRow(), self.ui.tb_Audience.currentColumn()):
            #self.ui.pb_Edit.setEnabled(True)
        #else:
            #self.ui.pb_Edit.setEnabled(False)

    def closeEvent(self,event):
        self.MainAppWindowShow=MainAppWindow()
        self.MainAppWindowShow.show()
        self.close()

    #Валидатор полей справочника аудиторий
    def validation(self):
        self.AudValid=Validator.AudienceValidator()

        self.AudName=self.ui.le_AudienceName.text()

        self.AudNaim=self.ui.tE_Naimen.toPlainText()
        self.AudTO=self.ui.tE_AudienceTO.toPlainText()
        self.AudPO=self.ui.tE_PO.toPlainText()

        if (self.AudValid.AudNameValid(self.AudName))==True:
            if (self.AudValid.AudNaimenValid(self.AudNaim))==True:
                if (self.AudValid.AudTOValid(self.AudTO))==True:
                    if (self.AudValid.AudPOValid(self.AudPO))==True:
                        return True
                    else:
                        self.APDialogUi.show()
                        return False
                else:
                     self.ATDialogUi.show()
                     return False
            else:
                self.ANMDialogUi.show()
                return False
        else:
            self.ANDialogUi.show()
            return False
        
        
    #Добавить запись
    def addRecord(self):
        if self.validation()==True:
            self.xName = str(self.ui.le_AudienceName.text())
            self.yName = str(self.ui.tE_AudienceTO.toPlainText())
            self.wName = str(self.ui.tE_Naimen.toPlainText())
            self.zName = str(self.ui.tE_PO.toPlainText())
            self.record = {'AudienceName': self.xName, 'AudiencePO' : self.yName, 'AudienceTO': self.wName, 'AudienceNaimenovanie' : self.zName }
            self.records.append(self.record)
            if len(self.records)>1:
                SelSortAud(self.records)
            writeCSV("AUDDB.csv",self.records)
            self.tableRecords()



    def tableRecords(self):
        self.records=AUDreadCSV("AUDDB.csv")
        if self.records:
            self.ui.tb_Audience.setRowCount(0)
            self.index = len(self.records)
            for i in range(0, self.index):
                self.rowCount= i
                self.ui.tb_Audience.insertRow(self.rowCount)
                self.ui.tb_Audience.setItem(self.rowCount, 0, QtWidgets.QTableWidgetItem(self.records[self.rowCount].get('AudienceName')))
                self.ui.tb_Audience.setItem(self.rowCount, 1, QtWidgets.QTableWidgetItem(self.records[self.rowCount].get('AudienceNaimenovanie')))
                self.ui.tb_Audience.setItem(self.rowCount, 2, QtWidgets.QTableWidgetItem(self.records[self.rowCount].get('AudienceTO')))
                self.ui.tb_Audience.setItem(self.rowCount, 3, QtWidgets.QTableWidgetItem(self.records[self.rowCount].get('AudiencePO')))
        self.ui.tb_Audience.setCurrentCell(0,0)

        
    def delRecord(self):
        self.ui.tb_Audience.removeRow(self.ui.tb_Audience.currentRow())
        self.row=self.ui.tb_Audience.currentRow()  
        self.ui.tb_Audience.setCurrentCell(self.row-1,0)
        self.records.pop(self.row)
        writeCSV("AUDDB.csv",self.records)
        
    def editRecord(self):
        self.ui.pb_Save.setEnabled(True)
        self.ui.pb_Add.setEnabled(False)
        self.ui.pb_Delete.setEnabled(False)
        self.ui.pb_Edit.setEnabled(False)
        

    def saveRecord(self):
        self.delRecord()
        self.addRecord()
        self.ui.pb_Save.setEnabled(False)
        self.ui.pb_Delete.setEnabled(True)
        self.ui.pb_Add.setEnabled(True)
        self.ui.pb_Edit.setEnabled(True)

    def ShowRecord(self,row,column):
        self.ui.le_AudienceName.setText(self.records[row].get("AudienceName"))
        self.ui.tE_AudienceTO.setPlainText(self.records[row].get("AudienceTO"))
        self.ui.tE_Naimen.setPlainText(self.records[row].get("AudienceNaimenovanie"))
        self.ui.tE_PO.setPlainText(self.records[row].get("AudiencePO"))


#Класс Редактирования КО
class KOEditorWindow(QtWidgets.QMainWindow):
    def __init__(self,):
        super(KOEditorWindow, self).__init__()
        self.ui=Ui_KO_Editor()
        self.counter=0
        self.records = []
        self.record={}
        self.index = 0
        self.disciplines=[]
        self.disc=''
        self.ui.setupUi(self)
        
        self.tableRecords()
        self.ui.pb_Save.clicked.connect(self.saveRecord)
        self.ui.pb_Add.clicked.connect(self.addRecord)
        self.ui.pb_Delete.clicked.connect(self.delRecord)
        self.ui.pb_Edit.clicked.connect(self.editRecord)
        self.ui.tb_KO.currentCellChanged.connect(self.ShowRecord)
        self.ui.pb_Save.setEnabled(False)

        self.FIODialogUi=FIODialog()
        self.FIODialogUi.setupUi(self)

        self.UslPrDialogUI=UslPrDialog()
        self.UslPrDialogUI.setupUi(self)

        self.NaprPodgotovDialogUi=NaprPodgotovDialog()
        self.NaprPodgotovDialogUi.setupUi(self)

        self.EducationDialogUi=EducationDialog()
        self.EducationDialogUi.setupUi(self)

    def closeEvent(self,event):
        self.MainAppWindowShow=MainAppWindow()
        self.MainAppWindowShow.show()
        self.close()

    def addRecord(self):
        if self.validation()==True:
            if self.ui.chB_State.isChecked()==False and self.ui.chB_Inner.isChecked()==False and self.ui.chB_Deal.isChecked()==False:
                self.UslPrDialogUI.show()
                return False
            else:
                self.fPPS = str(self.ui.le_FIO.text())
            if self.ui.chB_State.isChecked()==True:
                self.state = 1
            else:
                self.state = 0
            if self.ui.chB_Inner.isChecked()==True:
                self.inner = 1
            else:
                self.inner = 0
            if self.ui.chB_Deal.isChecked()==True:
                self.deal = 1
            else:
                self.deal = 0

            self.row=self.ui.tb_KO.currentRow()
            self.Dol = self.ui.cb_Dolzh.currentIndex()
            self.dRank = str(self.ui.cb_Dolzh.currentText())
            self.Step = self.ui.cb_Stepen.currentIndex()
            self.ucRank = str(self.ui.cb_Stepen.currentText())
            self.Zvan = self.ui.cb_zvan.currentIndex()
            self.zRank = str(self.ui.cb_zvan.currentText())
            self.nPPS = str(self.ui.tE_NaprPodgotov.toPlainText())
            self.ePPS = str(self.ui.tE_Education.toPlainText())
            self.record = {'FIO': self.fPPS,'Uslovia': [self.state,self.inner,self.deal], "Dolzhnost": self.Dol, "Stepen": self.Step, "Zvanie": self.Zvan, 'Napravlenie': self.nPPS, 'Education' : self.ePPS }
            self.records.append(self.record)
            if len(self.records)>1:
                SelSortPPS(self.records)
            writeCSV("PPSDB.csv",self.records)
            self.tableRecords()


    def tableRecords(self):
        self.records=PPSreadCSV("PPSDB.csv")
        if self.records:
            self.ui.tb_KO.setRowCount(0)
            self.index = len(self.records)

            for i in range(0, self.index):
                self.rowCount= i
                self.ui.tb_KO.insertRow(self.rowCount)
                if self.records[self.rowCount].get("Uslovia")[0]==1:     
                    self.c1PPS = str(self.ui.chB_State.text())
                else:
                    self.c1PPS  = str('')
                if self.records[self.rowCount].get("Uslovia")[1]==1:
                    self.c2PPS = str(self.ui.chB_Inner.text())
                else:
                    self.c2PPS  = str('')
                if self.records[self.rowCount].get("Uslovia")[2]==1:
                    self.c3PPS = str(self.ui.chB_Deal.text())
                else:
                    self.c3PPS  = str('')

                self.qboxPPS = self.c1PPS + ' ' + self.c2PPS + ' ' + self.c3PPS
                

                self.ui.tb_KO.setItem(self.rowCount, 0, QtWidgets.QTableWidgetItem(self.records[self.rowCount].get('FIO')))
                self.ui.tb_KO.setItem(self.rowCount, 1, QtWidgets.QTableWidgetItem(self.qboxPPS))
                self.ui.tb_KO.setItem(self.rowCount, 2, QtWidgets.QTableWidgetItem(self.ui.cb_Dolzh.itemText(self.records[self.rowCount].get("Dolzhnost"))))
                self.ui.tb_KO.setItem(self.rowCount, 3, QtWidgets.QTableWidgetItem(self.ui.cb_Stepen.itemText(self.records[self.rowCount].get("Stepen"))))
                self.ui.tb_KO.setItem(self.rowCount, 4, QtWidgets.QTableWidgetItem(self.ui.cb_zvan.itemText(self.records[self.rowCount].get("Zvanie"))))
                self.ui.tb_KO.setItem(self.rowCount, 5, QtWidgets.QTableWidgetItem(self.records[self.rowCount].get('Napravlenie')))
                self.ui.tb_KO.setItem(self.rowCount, 6, QtWidgets.QTableWidgetItem(self.records[self.rowCount].get('Education')))
            self.ui.tb_KO.setCurrentCell(0,0)


    def delRecord(self):
        self.ui.tb_KO.removeRow(self.ui.tb_KO.currentRow())
        self.row=self.ui.tb_KO.currentRow()
        self.records.pop(self.row)
        writeCSV("PPSDB.csv",self.records)
        
    def editRecord(self):
        self.ui.pb_Save.setEnabled(True)
        self.ui.pb_Add.setEnabled(False)
        self.ui.pb_Delete.setEnabled(False)
        self.ui.pb_Edit.setEnabled(False)
        

    def saveRecord(self):
        self.delRecord()
        self.addRecord()
        self.ui.pb_Save.setEnabled(False)
        self.ui.pb_Delete.setEnabled(True)
        self.ui.pb_Add.setEnabled(True)
        self.ui.pb_Edit.setEnabled(True)

    def ShowRecord(self,row,column):
        self.ui.le_FIO.setText(self.records[row].get("FIO"))
        if self.records[row].get("Uslovia")[0]==True:
            self.ui.chB_State.setChecked(True)
        else:
            self.ui.chB_State.setChecked(False)
        if self.records[row].get("Uslovia")[1]==True:
            self.ui.chB_Inner.setChecked(True)
        else:
            self.ui.chB_Inner.setChecked(False)
        if self.records[row].get("Uslovia")[2]==True:
            self.ui.chB_Deal.setChecked(True)
        else:
            self.ui.chB_Deal.setChecked(False)
        self.ui.cb_Dolzh.setCurrentIndex(self.records[row].get("Dolzhnost"))
        self.ui.cb_Stepen.setCurrentIndex(self.records[row].get("Stepen"))
        self.ui.cb_zvan.setCurrentIndex(self.records[row].get("Zvanie"))
        self.ui.tE_NaprPodgotov.setPlainText(self.records[row].get("Napravlenie"))
        self.ui.tE_Education.setPlainText(self.records[row].get("Education"))


    #Валидатор полей справочника аудиторий
    def validation(self):
        self.PPSValid=Validator.PPSValidator()

        self.FIO=self.ui.le_FIO.text()

        self.NaprPodgotov=self.ui.tE_NaprPodgotov.toPlainText()
        self.Education=self.ui.tE_NaprPodgotov.toPlainText()

        if (self.PPSValid.FIOValid(self.FIO))==True:
            return True
        else:
            self.FIODialogUi.show()
            return False

#Класс редактирования УП
class UPEditorWindow(QtWidgets.QMainWindow):
    def __init__(self):
        super(UPEditorWindow, self).__init__()
        self.ui=Ui_UPEditor()
        self.Audience=[]
        self.Teacher=[]
        self.TeachersAmount=1

        self.ui.setupUi(self)
        
        self.ui.TEACHER2.setVisible(False)
        self.ui.TEACHER3.setVisible(False)
        self.ui.TEACHER4.setVisible(False)
        self.ui.TEACHER5.setVisible(False)
        self.ui.TEACHER6.setVisible(False)

        self.ui.pb_Add2.clicked.connect(self.addTEACHER2)
        self.ui.pb_Add3.clicked.connect(self.addTEACHER3)
        self.ui.pb_Add4.clicked.connect(self.addTEACHER4)
        self.ui.pb_Add5.clicked.connect(self.addTEACHER5)
        self.ui.pb_Add6.clicked.connect(self.addTEACHER6)

        self.ui.pb_Del2.clicked.connect(self.remTEACHER2)
        self.ui.pb_Del3.clicked.connect(self.remTEACHER3)
        self.ui.pb_Del4.clicked.connect(self.remTEACHER4)
        self.ui.pb_Del5.clicked.connect(self.remTEACHER5)
        self.ui.pb_Del6.clicked.connect(self.remTEACHER6)

        self.ui.pb_Save.clicked.connect(self.saveRecord)
        self.ui.pb_Add.clicked.connect(self.addRecord)
        self.ui.pb_Delete.clicked.connect(self.delRecord)
        self.ui.pb_Edit.clicked.connect(self.editRecord)

        self.ui.list_AllAud.itemDoubleClicked.connect(self.addAud)
        self.ui.list_ChosAud.itemDoubleClicked.connect(self.removeAud)
        self.ui.list_Disc.currentRowChanged.connect(self.ShowRecord)

        self.ui.pb_Save.setEnabled(False)

        self.tableRecords()
        self.LoadKOAndAud()


    def addTEACHER2(self):
        self.ui.TEACHER2.setVisible(True)
        self.TeachersAmount=2
    
    def addTEACHER3(self):
        self.ui.TEACHER3.setVisible(True)
        self.TeachersAmount=3

    def addTEACHER4(self):
        self.ui.TEACHER4.setVisible(True)
        self.TeachersAmount=4

    def addTEACHER5(self):
        self.ui.TEACHER5.setVisible(True)
        self.TeachersAmount=5

    def addTEACHER6(self):
        self.ui.TEACHER6.setVisible(True)
        self.TeachersAmount=6



    def remTEACHER2(self):
        if self.TeachersAmount==6:
            self.ui.cb_Teacher2.setCurrentIndex(self.ui.cb_Teacher3.currentIndex())
            self.ui.cb_Teacher3.setCurrentIndex(self.ui.cb_Teacher4.currentIndex())
            self.ui.cb_Teacher4.setCurrentIndex(self.ui.cb_Teacher5.currentIndex())
            self.ui.cb_Teacher5.setCurrentIndex(self.ui.cb_Teacher6.currentIndex())
            self.ui.cb_Teacher6.setCurrentIndex(0)
            self.ui.TEACHER6.setVisible(False)
        elif self.TeachersAmount==5:
            self.ui.cb_Teacher2.setCurrentIndex(self.ui.cb_Teacher3.currentIndex())
            self.ui.cb_Teacher3.setCurrentIndex(self.ui.cb_Teacher4.currentIndex())
            self.ui.cb_Teacher4.setCurrentIndex(self.ui.cb_Teacher5.currentIndex())
            self.ui.cb_Teacher5.setCurrentIndex(0)
            self.ui.TEACHER5.setVisible(False)
        elif self.TeachersAmount==4:
            self.ui.cb_Teacher2.setCurrentIndex(self.ui.cb_Teacher3.currentIndex())
            self.ui.cb_Teacher3.setCurrentIndex(self.ui.cb_Teacher4.currentIndex())
            self.ui.cb_Teacher4.setCurrentIndex(0)
            self.ui.TEACHER4.setVisible(False)
        elif self.TeachersAmount==3:
            self.ui.cb_Teacher2.setCurrentIndex(self.ui.cb_Teacher3.currentIndex())
            self.ui.cb_Teacher3.setCurrentIndex(0)
            self.ui.TEACHER3.setVisible(False)
        elif self.TeachersAmount==2:
            self.ui.cb_Teacher2.setCurrentIndex(0)
            self.ui.TEACHER2.setVisible(False)
        self.TeachersAmount=self.TeachersAmount-1
        print(self.TeachersAmount)
    
    def remTEACHER3(self):
        if self.TeachersAmount==6:
            self.ui.cb_Teacher3.setCurrentIndex(self.ui.cb_Teacher4.currentIndex())
            self.ui.cb_Teacher4.setCurrentIndex(self.ui.cb_Teacher5.currentIndex())
            self.ui.cb_Teacher5.setCurrentIndex(self.ui.cb_Teacher6.currentIndex())
            self.ui.cb_Teacher6.setCurrentIndex(0)
            self.ui.TEACHER6.setVisible(False)
        elif self.TeachersAmount==5:
            self.ui.cb_Teacher3.setCurrentIndex(self.ui.cb_Teacher4.currentIndex())
            self.ui.cb_Teacher4.setCurrentIndex(self.ui.cb_Teacher5.currentIndex())
            self.ui.cb_Teacher5.setCurrentIndex(0)
            self.ui.TEACHER5.setVisible(False)
        elif self.TeachersAmount==4:
            self.ui.cb_Teacher3.setCurrentIndex(self.ui.cb_Teacher4.currentIndex())
            self.ui.cb_Teacher4.setCurrentIndex(0)
            self.ui.TEACHER4.setVisible(False)
        elif self.TeachersAmount==3:
            self.ui.cb_Teacher3.setCurrentIndex(0)
            self.ui.TEACHER3.setVisible(False)
        self.TeachersAmount=self.TeachersAmount-1

    def remTEACHER4(self):
        if self.TeachersAmount==6:
            self.ui.cb_Teacher4.setCurrentIndex(self.ui.cb_Teacher5.currentIndex())
            self.ui.cb_Teacher5.setCurrentIndex(self.ui.cb_Teacher6.currentIndex())
            self.ui.cb_Teacher6.setCurrentIndex(0)
            self.ui.TEACHER6.setVisible(False)
        elif self.TeachersAmount==5:
            self.ui.cb_Teacher4.setCurrentIndex(self.ui.cb_Teacher5.currentIndex())
            self.ui.cb_Teacher5.setCurrentIndex(0)
            self.ui.TEACHER5.setVisible(False)
        elif self.TeachersAmount==4:
            self.ui.cb_Teacher4.setCurrentIndex(0)
            self.ui.TEACHER4.setVisible(False)
        self.TeachersAmount=self.TeachersAmount-1

    def remTEACHER5(self):
        if self.TeachersAmount==6:
            self.ui.cb_Teacher5.setCurrentIndex(self.ui.cb_Teacher6.currentIndex())
            self.ui.cb_Teacher6.setCurrentIndex(0)
            self.ui.TEACHER6.setVisible(False)
        elif self.TeachersAmount==5:
            self.ui.cb_Teacher5.setCurrentIndex(0)
            self.ui.TEACHER5.setVisible(False)
        self.TeachersAmount=self.TeachersAmount-1

    def remTEACHER6(self):
        self.ui.cb_Teacher6.setCurrentIndex(0)
        self.ui.TEACHER6.setVisible(False)
        self.TeachersAmount=self.TeachersAmount-1


    def addAud(self,item):
        if len(self.ui.list_ChosAud.findItems(item.text(),QtCore.Qt.MatchExactly))<1:
            self.ui.list_ChosAud.addItem(item.text())

    def removeAud(self,item):
        self.ui.list_ChosAud.takeItem(self.ui.list_ChosAud.currentRow())

    def closeEvent(self,event):
        self.MainAppWindowShow=MainAppWindow()
        self.MainAppWindowShow.show()
        self.close()

    def LoadKOAndAud(self):
        self.Teachers=TeacherreadCSV("PPSDB.csv")
        if self.Teachers:
            for j in self.Teachers:
                self.ui.cb_Teacher1.addItem(j)
                self.ui.cb_Teacher2.addItem(j)
                self.ui.cb_Teacher3.addItem(j)
                self.ui.cb_Teacher4.addItem(j)
                self.ui.cb_Teacher5.addItem(j)
                self.ui.cb_Teacher6.addItem(j)
        self.Audiences=TeacherreadCSV("AUDDB.csv")
        if self.Audiences:
            for j in range(len(self.Audiences)):
                self.ui.list_AllAud.addItem(self.Audiences[j])


    def addRecord(self):
        if True==True:
            self.NameUD = str(self.ui.le_NameUD.text())
            self.NumberUD = str(self.ui.le_NumberUD.text()) 
            for i in range(0,self.ui.list_ChosAud.count()):
                self.Audience.append(self.ui.list_ChosAud.item(i).text())

            if self.TeachersAmount==1:
                self.Teacher.append(self.ui.cb_Teacher1.currentText())

            elif self.TeachersAmount==2:
                self.Teacher.append(self.ui.cb_Teacher1.currentText())
                self.Teacher.append(self.ui.cb_Teacher2.currentText())

            elif self.TeachersAmount==3:
                self.Teacher.append(self.ui.cb_Teacher1.currentText())
                self.Teacher.append(self.ui.cb_Teacher2.currentText())
                self.Teacher.append(self.ui.cb_Teacher3.currentText())

            elif self.TeachersAmount==4:
                self.Teacher.append(self.ui.cb_Teacher1.currentText())
                self.Teacher.append(self.ui.cb_Teacher2.currentText())
                self.Teacher.append(self.ui.cb_Teacher3.currentText())
                self.Teacher.append(self.ui.cb_Teacher4.currentText())

            elif self.TeachersAmount==5:
                self.Teacher.append(self.ui.cb_Teacher1.currentText())
                self.Teacher.append(self.ui.cb_Teacher2.currentText())
                self.Teacher.append(self.ui.cb_Teacher3.currentText())
                self.Teacher.append(self.ui.cb_Teacher4.currentText())
                self.Teacher.append(self.ui.cb_Teacher5.currentText())

            elif self.TeachersAmount==6:
                self.Teacher.append(self.ui.cb_Teacher1.currentText())
                self.Teacher.append(self.ui.cb_Teacher2.currentText())
                self.Teacher.append(self.ui.cb_Teacher3.currentText())
                self.Teacher.append(self.ui.cb_Teacher4.currentText())
                self.Teacher.append(self.ui.cb_Teacher5.currentText())
                self.Teacher.append(self.ui.cb_Teacher6.currentText())

            self.record = {'NameUD': self.NameUD, 'NumberUD' : self.NumberUD, 'Teacher': self.Teacher, "Audience":self.Audience, "Amount":self.TeachersAmount }
            print(self.record)
            self.records.append(self.record)
            writeCSV("UPDB.csv",self.records)
            self.Teacher.clear()
            self.TeachersAmount=1
            self.tableRecords()

    def tableRecords(self):
        self.ui.list_Disc.clear()
        self.records=UPreadCSV("UPDB.csv")
        if self.records:
            for rec in self.records:
                self.ui.list_Disc.addItem(rec.get("NameUD"))
        self.ui.list_Disc.setCurrentRow(0)
        

    def delRecord(self):
        self.ui.list_Disc.removeItemWidget(self.ui.list_Disc.currentItem())
        if self.ui.list_Disc.currentRow()!=1:
            self.ui.list_Disc.setCurrentRow(self.ui.list_Disc.currentRow()-1)
        elif self.ui.list_Disc.count()>self.ui.list_Disc.currentRow():
            self.ui.list_Disc.setCurrentRow(self.ui.list_Disc.currentRow())
        self.records.pop(self.ui.list_Disc.currentRow())
        writeCSV("UPDB.csv",self.records)
        self.tableRecords()
        
        
    def editRecord(self):
        self.ui.pb_Save.setEnabled(True)
        self.ui.pb_Add.setEnabled(False)
        self.ui.pb_Delete.setEnabled(False)
        self.ui.pb_Edit.setEnabled(False)
        

    def saveRecord(self):
        self.delRecord()
        self.addRecord()
        self.ui.pb_Save.setEnabled(False)
        self.ui.pb_Delete.setEnabled(True)
        self.ui.pb_Add.setEnabled(True)
        self.ui.pb_Edit.setEnabled(True)

    def ShowRecord(self,item):
        self.ui.list_ChosAud.clear()
        temp = re.findall(r'[А-я]+\-\d{3}',self.records[self.ui.list_Disc.currentRow()].get("Audience"))
        res = list(map(str, temp))
        for i in range(len(res)):
            self.ui.list_ChosAud.addItem(res[i])
        temp = re.findall(r'([А-я]+\ [А-я]+\ [А-я]+)',self.records[self.ui.list_Disc.currentRow()].get("Teacher"))
        res = list(map(str, temp))

        self.ui.le_NameUD.setText(self.records[self.ui.list_Disc.currentRow()].get("NameUD"))
        self.ui.le_NumberUD.setText(self.records[self.ui.list_Disc.currentRow()].get("NumberUD"))

        self.ui.TEACHER2.setVisible(False)
        self.ui.TEACHER3.setVisible(False)
        self.ui.TEACHER4.setVisible(False)
        self.ui.TEACHER5.setVisible(False)
        self.ui.TEACHER6.setVisible(False)

        self.TeachersAmount=int(self.records[self.ui.list_Disc.currentRow()].get("Amount"))
        print(self.records[self.ui.list_Disc.currentRow()])

        if self.records[self.ui.list_Disc.currentRow()].get("Amount")==str(1):
            self.ui.cb_Teacher1.setCurrentIndex(self.ui.cb_Teacher1.findText(res[0]))
        elif self.records[self.ui.list_Disc.currentRow()].get("Amount")==str(2):
            self.ui.cb_Teacher1.setCurrentIndex(self.ui.cb_Teacher1.findText(res[0]))
            self.ui.cb_Teacher2.setCurrentIndex(self.ui.cb_Teacher2.findText(res[1]))
            self.ui.TEACHER2.setVisible(True)
        elif self.records[self.ui.list_Disc.currentRow()].get("Amount")==str(3):
            self.ui.cb_Teacher1.setCurrentIndex(self.ui.cb_Teacher1.findText(res[0]))
            self.ui.cb_Teacher2.setCurrentIndex(self.ui.cb_Teacher2.findText(res[1]))
            self.ui.cb_Teacher3.setCurrentIndex(self.ui.cb_Teacher3.findText(res[2]))
            self.ui.TEACHER2.setVisible(True)
            self.ui.TEACHER3.setVisible(True)
        elif self.records[self.ui.list_Disc.currentRow()].get("Amount")==str(4):
            self.ui.cb_Teacher1.setCurrentIndex(self.ui.cb_Teacher1.findText(res[0]))
            self.ui.cb_Teacher2.setCurrentIndex(self.ui.cb_Teacher2.findText(res[1]))
            self.ui.cb_Teacher3.setCurrentIndex(self.ui.cb_Teacher3.findText(res[2]))
            self.ui.cb_Teacher4.setCurrentIndex(self.ui.cb_Teacher4.findText(res[3]))
            self.ui.TEACHER2.setVisible(True)
            self.ui.TEACHER3.setVisible(True)
            self.ui.TEACHER4.setVisible(True)
        elif self.records[self.ui.list_Disc.currentRow()].get("Amount")==str(5):
            self.ui.cb_Teacher1.setCurrentIndex(self.ui.cb_Teacher1.findText(res[0]))
            self.ui.cb_Teacher2.setCurrentIndex(self.ui.cb_Teacher2.findText(res[1]))
            self.ui.cb_Teacher3.setCurrentIndex(self.ui.cb_Teacher3.findText(res[2]))
            self.ui.cb_Teacher4.setCurrentIndex(self.ui.cb_Teacher4.findText(res[3]))
            self.ui.cb_Teacher5.setCurrentIndex(self.ui.cb_Teacher5.findText(res[4]))
            self.ui.TEACHER2.setVisible(True)
            self.ui.TEACHER3.setVisible(True)
            self.ui.TEACHER4.setVisible(True)
            self.ui.TEACHER5.setVisible(True)
        elif self.records[self.ui.list_Disc.currentRow()].get("Amount")==str(6):
            self.ui.cb_Teacher1.setCurrentIndex(self.ui.cb_Teacher1.findText(res[0]))
            self.ui.cb_Teacher2.setCurrentIndex(self.ui.cb_Teacher2.findText(res[1]))
            self.ui.cb_Teacher3.setCurrentIndex(self.ui.cb_Teacher3.findText(res[2]))
            self.ui.cb_Teacher4.setCurrentIndex(self.ui.cb_Teacher4.findText(res[3]))
            self.ui.cb_Teacher5.setCurrentIndex(self.ui.cb_Teacher5.findText(res[4]))
            self.ui.cb_Teacher6.setCurrentIndex(self.ui.cb_Teacher6.findText(res[5]))
            self.ui.TEACHER2.setVisible(True)
            self.ui.TEACHER3.setVisible(True)
            self.ui.TEACHER4.setVisible(True)
            self.ui.TEACHER5.setVisible(True)
            self.ui.TEACHER6.setVisible(True)

    """def validation(self):
        self.UPValid=Validator.UPValidator()

        self.Name=self.ui.le_NameUD.text()
        self.Number=self.ui.le_NumberUD.text()

        if (self.UPValid.NameValid(self.Name))==True:
            if (self.UPValid.NumberValid(self.Number))==True:
                if self.ui.lt_ChosenTeacher.count()!=0:
                    if self.ui.lt_ChosenAudience.count()!=0:
                        return True
                    else:
                        self.AudDialogUi.show()
                        return False
                else:
                    self.TeacherDialogUi.show()
                    return False
            else:
                self.NumberUDDialogUi.show()
                return False
        else:
            self.NameUDDialogUi.show()
            return False"""

#Класс главного окна
class MainAppWindow(QtWidgets.QMainWindow):
    def __init__(self):
        super(MainAppWindow, self).__init__()
        self.ui=Ui_MainWindow()
        self.Document=-1
        self.DocRecords=[]
        self.ui.setupUi(self)
        self.setWindowTitle("Приложение для генерации справок")
        #self.ui.tb_AvailAud.setVisible(False)
        #self.ui.tb_ChosenAdu.setVisible(False)
        #self.ui.tb_AvailTeacher.setVisible(False)
        #self.ui.tb_ChosenTeacher.setVisible(False)

        self.ui.pushButton.clicked.connect(self.save)
        self.ui.action_TemplateMTO.triggered.connect(self.MTOChosen)
        self.ui.action_TemplateKO.triggered.connect(self.KOChosen)

        #Вызов файлового менеджера
        self.ui.action_Open.triggered.connect(self.openFile)
        #Создание файла
        self.ui.action_Save.triggered.connect(self.saveFile)

        """Кнопки вызова дочерних окон"""
        self.ui.action_PPS.triggered.connect(self.openKOEditor)
        self.ui.action_GN.triggered.connect(self.openUPEditor)
        self.ui.action_Audience.triggered.connect(self.openAudienceEditor)

    def save(self):
        if self.Document==1:
            self.DocRecords=matchingMTO(readMTODisc("UPDB.csv"),readMTOAUD("AUDDB.csv"))
            self.doc = docx.Document('Testooo.docx')
            self.table = self.doc.add_table(rows=1,cols=5, style='Table Grid')
            hdr_cells = self.table.rows[0].cells
            hdr_cells[0].text = '№'
            hdr_cells[1].text = 'Наименование дисциплины (модуля) практик в соответствии с учебным планом'
            hdr_cells[2].text = 'Наименование специальных помещений и помещений для самостоятельной работы'
            hdr_cells[3].text = 'Оснащенность специальных помещений и помещений для самостоятельной работы'
            hdr_cells[4].text = 'Перечень лицензионного программного обеспечения. Реквизиты подтверждающего документа'
            for i in self.DocRecords:
                row_cells = self.table.add_row().cells
                row_cells[1].text = i.get('Discipline')
                row_cells[2].text = i.get('AudiencePO')
                row_cells[3].text = i.get('AudienceTO')
                row_cells[4].text = i.get('AudienceNaimenovanie')
            self.filename=QtWidgets.QFileDialog.getSaveFileName(self, "Выберите файл", os.getcwd(), ".DOCX Файлы (*.docx)")
            directory=str(self.filename)
            cleanDirectory=""
            counter=0
            for i  in range(2,len(directory)):
                if directory[i]=="\'":
                    counter+=1
                if counter<1:
                    cleanDirectory=cleanDirectory+directory[i]
            self.doc.save(cleanDirectory)

            
        elif self.Document==2:
            self.DocRecords=matchingKO(readKODisc("UPDB.csv"),readKOTeacher("PPSDB.csv"))
            self.doc = docx.Document('Testoko.docx')
            self.table = self.doc.add_table(rows=1,cols=7, style='Table Grid')
            hdr_cells = self.table.rows[0].cells
            hdr_cells[0].text = '№'
            hdr_cells[1].text = 'Ф.И.О. преподавателя, реализующего программу '
            hdr_cells[2].text = 'Условия привлечения (основное место работы: штатный, внутренний совместитель, внешний совместитель по договору ГПХ)'
            hdr_cells[3].text = 'Должность, ученая степень, ученое звание'
            hdr_cells[4].text = 'Перечень читаемых дисциплин '
            hdr_cells[5].text = 'Уровень образования,наименование специальности,направления подготовки,наименование присвоенной квалификации'
            hdr_cells[6].text = 'Сведения о дополнительном профессиональном образовании '
            for i in self.DocRecords:
                row_cells = self.table.add_row().cells
                row_cells[1].text = i.get('FIO')
                row_cells[2].text = str(i.get('Uslovia'))
                row_cells[3].text = str(i.get('Dolzhnost'))+str(i.get('Stepen'))+str(i.get('Zvanie'))
                row_cells[4].text = i.get('Discipline')
                row_cells[5].text = i.get('Napravlenie')
                row_cells[6].text = str(i.get('Education'))

            self.filename=QtWidgets.QFileDialog.getSaveFileName(self, "Выберите файл", os.getcwd(), ".DOCX Файлы (*.docx)")
            directory=str(self.filename)
            cleanDirectory=""
            counter=0
            for i  in range(2,len(directory)):
                if directory[i]=="\'":
                    counter+=1
                if counter<1:
                    cleanDirectory=cleanDirectory+directory[i]
            self.doc.save(cleanDirectory)



    def MTOChosen(self):
        self.Document=1

    def KOChosen(self):
        self.Document=2


    #Функция вызова файлового менеджера
    def openFile(self):
        self.filenameDocx = QtWidgets.QFileDialog.getOpenFileName(self, "Выберите файл", os.getcwd(), ".DOCX Файлы (*.docx)")
        directory=str(self.filenameDocx)
        docxDirectory=""
        counter=0
        for i  in range(2,len(directory)):
            if directory[i]=="\'":
                counter+=1
            if counter<1:
                docxDirectory=docxDirectory+directory[i]

        self.filenamePDF = QtWidgets.QFileDialog.getSaveFileName(self, "Выберите файл", os.getcwd(), ".pdf Файлы (*.pdf)")
        directory=str(self.filenamePDF)
        PDFDirectory=""
        counter=0
        for i  in range(2,len(directory)):
            if directory[i]=="\'":
                counter+=1
            if counter<1:
                PDFDirectory=PDFDirectory+directory[i]

        convertDocxToPDF.convertDocxToPdf(docxDirectory,PDFDirectory)

    def saveFile(self):
        self.filename=QtWidgets.QFileDialog.getSaveFileName(self, "Выберите файл", os.getcwd(), ".DOCX Файлы (*.docx)")
        directory=str(self.filename)
        cleanDirectory=""
        counter=0
        for i  in range(2,len(directory)):
            if directory[i]=="\'":
                counter+=1
            if counter<1:
                cleanDirectory=cleanDirectory+directory[i]

        DocxGeneratingDef.allInOne(cleanDirectory)

    """Функции вызова дочерних окон"""

    def openKOEditor(self):
        self.KOEditor=KOEditorWindow()
        self.KOEditor.show()
        self.hide()

    def openUPEditor(self):
        self.UPEditingWindow=UPEditorWindow()
        self.UPEditingWindow.show()
        self.hide()

    def openAudienceEditor(self):
        self.AudienceEditingWindow=AudienceEditorWindow()
        self.AudienceEditingWindow.show()
        self.hide()



           # self.doc = docx.Document('Testooo.docx')
            #self.df = pd.DataFrame()
            #with open('AUDDB.csv', newline='') as File:
             #   reader = csv.reader(File)
              #  headers = next(reader)
               # cols = len(headers)
                #for self.table in self.doc.tables:
                #    hdr_cells = self.table.rows[0].cells
                #for i in range(cols):
                #    hdr_cells[i].text = headers[i]
                #for row in reader:
                #    row_cells = self.table.add_row().cells
                #    for i in range(cols):
                 #       row_cells[i].text = row[i]
                #self.doc.save('Testooo.docx')

  
  
    



#Вызов программы
if __name__ == "__main__":
    app=QtWidgets.QApplication(sys.argv)
    MainWindow=MainAppWindow()
    MainWindow.show()
    sys.exit(app.exec_())


